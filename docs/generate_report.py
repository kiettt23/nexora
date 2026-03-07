# -*- coding: utf-8 -*-
"""Generate BaoCao.docx for Nexora project — v2 with visual polish."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Constants
NAVY = RGBColor(0, 42, 82)       # #002A52
ACCENT = RGBColor(0, 102, 178)   # #0066B2
DARK = RGBColor(30, 30, 30)
GRAY = RGBColor(100, 100, 100)
WHITE = RGBColor(255, 255, 255)
HEADER_BG = "002A52"
ACCENT_BG = "E8F0FE"
ROW_ALT_BG = "F5F7FA"

# ── Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

# ── Default style
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
font.color.rgb = DARK
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

# ══════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════

def set_cell_bg(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = NAVY if level <= 2 else ACCENT
    if level == 1:
        # Add bottom border line
        pPr = h._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="{HEADER_BG}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    return h

def add_para(text, bold=False, italic=False, align=None, spacing_after=6, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size or 13)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color or DARK
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    return p

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(13)
        r1.bold = True
        r1.font.color.rgb = DARK
        r2 = p.add_run(text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(13)
        r2.font.color.rgb = DARK
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.font.color.rgb = DARK
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def add_table_styled(headers, rows, col_widths=None):
    """Create a professional table with colored header and alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, HEADER_BG)

    # Data rows with alternating bg
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = DARK
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_idx % 2 == 1:
                set_cell_bg(cell, ROW_ALT_BG)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    # Table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()  # spacing after table
    return table

def add_code_block(text, font_size=9):
    """Add a monospaced code block with light background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Background shading for paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F2F5" w:val="clear"/>')
    pPr.append(shading)
    # Add border
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="2" w:color="D0D0D0"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="4" w:color="D0D0D0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="2" w:color="D0D0D0"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="4" w:color="D0D0D0"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(40, 40, 40)
    return p

def add_info_box(text, bg_color="E8F0FE", border_color="0066B2"):
    """Add a highlighted info box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
    pPr.append(shading)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="284"/>')
    pPr.append(ind)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    return p

def add_page_number():
    """Add page number to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Page number field
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)
        for r in [run, run2, run3]:
            r.font.size = Pt(10)
            r.font.name = 'Times New Roman'
            r.font.color.rgb = GRAY

add_page_number()

# ══════════════════════════════════════════════════════════════════
# TRANG BÌA
# ══════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

add_para('TRƯỜNG ĐẠI HỌC FPT', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=NAVY, spacing_after=2)
add_para('BỘ MÔN LẬP TRÌNH C# 3', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, color=ACCENT, spacing_after=24)

for _ in range(3):
    doc.add_paragraph()

# Title with decorative line
line_before = doc.add_paragraph()
line_before.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = line_before._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="12" w:space="8" w:color="{HEADER_BG}"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(4)
run_t = p_title.add_run('BÁO CÁO ASSIGNMENT')
run_t.font.name = 'Times New Roman'
run_t.font.size = Pt(28)
run_t.bold = True
run_t.font.color.rgb = NAVY

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(4)
run_s = p_sub.add_run('Thiết kế Web bán hàng')
run_s.font.name = 'Times New Roman'
run_s.font.size = Pt(20)
run_s.font.color.rgb = ACCENT

p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_n = p_name.add_run('NEXORA')
run_n.font.name = 'Times New Roman'
run_n.font.size = Pt(36)
run_n.bold = True
run_n.font.color.rgb = NAVY

# Decorative line after
line_after = doc.add_paragraph()
line_after.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr2 = line_after._p.get_or_add_pPr()
pBdr2 = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="12" w:space="8" w:color="{HEADER_BG}"/>'
    f'</w:pBdr>'
)
pPr2.append(pBdr2)

for _ in range(3):
    doc.add_paragraph()

# Tech stack info box
tech_table = doc.add_table(rows=4, cols=2)
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
tech_info = [
    ('Nền tảng:', 'ASP.NET Core 8 MVC'),
    ('CSDL:', 'PostgreSQL + Entity Framework Core'),
    ('Giao diện:', 'Razor View + Tailwind CSS + DaisyUI'),
    ('Ngôn ngữ:', 'C# / HTML / CSS / JavaScript'),
]
for i, (label, val) in enumerate(tech_info):
    c1 = tech_table.rows[i].cells[0]
    c1.text = ''
    r1 = c1.paragraphs[0].add_run(label)
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.font.color.rgb = NAVY
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c1.width = Cm(4)

    c2 = tech_table.rows[i].cells[1]
    c2.text = ''
    r2 = c2.paragraphs[0].add_run(val)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.font.color.rgb = DARK
    c2.width = Cm(10)

# Remove table borders for cover page
tbl = tech_table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'</w:tblBorders>'
)
tblPr.append(borders)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# MỤC LỤC
# ══════════════════════════════════════════════════════════════════
add_heading_styled('MỤC LỤC', level=1)

toc_sections = [
    ('1.', 'Giới thiệu dự án'),
    ('2.', 'Thiết kế cơ sở dữ liệu (Y1)'),
    ('', '   2.1. Sơ đồ quan hệ (ERD)'),
    ('', '   2.2. Chi tiết các bảng'),
    ('', '   2.3. Giải thích cách thiết kế'),
    ('', '   2.4. Dữ liệu mẫu'),
    ('3.', 'Sơ đồ Use Case'),
    ('4.', 'Cấu trúc dự án ASP.NET Core MVC'),
    ('', '   4.1. Mô hình MVC'),
    ('', '   4.2. Cấu trúc thư mục'),
    ('', '   4.3. Công nghệ sử dụng'),
    ('5.', 'Các chức năng đã thực hiện'),
    ('', '   5.1. Chức năng khách hàng'),
    ('', '   5.2. Chức năng quản lý'),
    ('', '   5.3. Chức năng nâng cao'),
    ('6.', 'Bảo mật và phân quyền'),
    ('7.', 'Triển khai'),
]
for num, title in toc_sections:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if num:
        r1 = p.add_run(num + ' ')
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(13)
        r1.font.color.rgb = NAVY
        r2 = p.add_run(title)
        r2.bold = True
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(13)
        r2.font.color.rgb = DARK
    else:
        r = p.add_run(title)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = GRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. GIỚI THIỆU
# ══════════════════════════════════════════════════════════════════
add_heading_styled('1. Giới thiệu dự án', level=1)

add_info_box(
    'Nexora — Ứng dụng web thương mại điện tử chuyên bán sản phẩm công nghệ '
    '(điện thoại, laptop, máy tính bảng, đồng hồ thông minh, phụ kiện). '
    'Xây dựng trên ASP.NET Core 8 MVC + EF Core + PostgreSQL.'
)

add_para(
    'Ứng dụng phục vụ hai nhóm người dùng chính:'
)
add_bullet('Khách hàng: Duyệt sản phẩm, tìm kiếm, thêm giỏ hàng, đặt hàng, '
           'sử dụng mã giảm giá, quản lý đơn hàng và thông tin cá nhân.', bold_prefix='Khách hàng: ')
add_bullet('Quản trị: Quản lý sản phẩm (CRUD + upload ảnh Cloudinary), danh mục, '
           'đơn hàng, người dùng, phân quyền 3 cấp, cấu hình cửa hàng.', bold_prefix='Quản trị: ')

add_para('')
add_para('Điểm nổi bật:', bold=True, color=NAVY)
add_bullet('56 sản phẩm + 8 danh mục + 8 voucher được seed sẵn')
add_bullet('Upload ảnh qua Cloudinary CDN')
add_bullet('Hệ thống voucher giảm giá đa dạng (%, cố định, min order, max discount)')
add_bullet('Phân quyền 3 cấp: Admin > Staff > Customer')
add_bullet('Responsive trên mọi thiết bị (Tailwind CSS + DaisyUI)')
add_bullet('Triển khai production trên Railway')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 2. THIẾT KẾ CSDL
# ══════════════════════════════════════════════════════════════════
add_heading_styled('2. Thiết kế cơ sở dữ liệu (Y1)', level=1)

# ── 2.1 ERD
add_heading_styled('2.1. Sơ đồ quan hệ (ERD)', level=2)

add_para('Sơ đồ thể hiện quan hệ giữa các bảng trong CSDL:', italic=True, color=GRAY)

erd = """
  ┌──────────────────┐          ┌──────────────────┐
  │   Categories     │          │   ShopConfigs    │
  │──────────────────│          │──────────────────│
  │ Id (PK)          │          │ Id (PK)          │
  │ Name             │          │ Key (Unique)     │
  │ Slug (Unique)    │          │ Value            │
  │ Description      │          │ Type             │
  │ SortOrder        │          └──────────────────┘
  │ IsActive         │
  └────────┬─────────┘          ┌──────────────────┐
           │ 1:N                │    Vouchers      │
           ▼                    │──────────────────│
  ┌──────────────────┐          │ Id (PK)          │
  │    Products      │          │ Code (Unique)    │
  │──────────────────│          │ DiscountAmount   │
  │ Id (PK)          │          │ DiscountPercent  │
  │ Name             │          │ MinOrderAmount   │
  │ Slug (Unique)    │          │ MaxDiscountAmount│
  │ Price            │          │ UsageLimit       │
  │ Brand, Color     │          │ StartDate        │
  │ Storage, RAM     │          │ EndDate          │
  │ ScreenSize, CPU  │          │ IsActive         │
  │ CategoryId (FK)──┤          └──────────────────┘
  │ IsActive         │
  │ IsFeatured       │
  └──┬───────────┬───┘
     │ 1:N       │ N:1 (Restrict)
     ▼           │
  ┌──────────────┐    │    ┌──────────────────────┐
  │ProductImages │    │    │  AspNetUsers          │
  │──────────────│    │    │  (IdentityUser)       │
  │ Id (PK)      │    │    │──────────────────────│
  │ ProductId(FK)│    │    │ Id (PK)              │
  │ ImagePath    │    │    │ FullName, Phone      │
  │ SortOrder    │    │    │ Address, Avatar      │
  │ IsMain       │    │    │ IsActive, CreatedAt  │
  └──────────────┘    │    └───┬──────────┬───────┘
                      │        │ 1:1      │ 1:N
                      │        ▼          ▼
                      │  ┌──────────┐  ┌──────────────┐
                      │  │  Carts   │  │   Orders     │
                      │  │──────────│  │──────────────│
                      │  │ Id (PK)  │  │ Id (PK)      │
                      │  │UserId(FK)│  │ UserId (FK)  │
                      │  └────┬─────┘  │ OrderCode    │
                      │       │ 1:N    │ TotalAmount  │
                      │       ▼        │ VoucherCode  │
                      │ ┌────────────┐ │ Status       │
                      │ │CartDetails │ └──────┬───────┘
                      │ │────────────│        │ 1:N
                      │ │ Id (PK)    │        ▼
                      │ │ CartId(FK) │ ┌──────────────┐
                      └─│ProductId(FK│ │ OrderDetails │
                        │ Quantity   │ │──────────────│
                        │ UnitPrice  │ │ Id (PK)      │
                        └────────────┘ │ OrderId (FK) │
                                       │ ProductId(FK)│
                                       │ ProductName  │
                                       │ Quantity     │
                                       │ UnitPrice    │
                                       └──────────────┘"""

add_code_block(erd, font_size=8)

add_para('')
add_info_box(
    'Ghi chú: Quan hệ Cascade = xóa bản ghi cha sẽ tự động xóa con. '
    'Restrict = không cho xóa cha nếu còn bản ghi con tham chiếu.'
)

doc.add_page_break()

# ── 2.2 Chi tiết các bảng
add_heading_styled('2.2. Chi tiết các bảng', level=2)

# Products
add_para('BẢNG PRODUCTS', bold=True, color=NAVY, size=14)
add_table_styled(
    ['Cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['Id', 'int', 'PK, Identity', 'Khóa chính'],
        ['Name', 'nvarchar(200)', 'Required', 'Tên sản phẩm'],
        ['Slug', 'nvarchar(200)', 'Required, Unique', 'URL thân thiện SEO'],
        ['Description', 'nvarchar(5000)', 'Nullable', 'Mô tả chi tiết'],
        ['Price', 'decimal(18,0)', '', 'Giá bán (VNĐ)'],
        ['OriginalPrice', 'decimal(18,0)', 'Nullable', 'Giá gốc (hiện giảm giá)'],
        ['Brand', 'nvarchar(100)', 'Nullable', 'Thương hiệu (Samsung, Apple...)'],
        ['Color', 'nvarchar(50)', 'Nullable', 'Màu sắc'],
        ['Storage', 'nvarchar(50)', 'Nullable', 'Dung lượng (128GB, 256GB...)'],
        ['RAM', 'nvarchar(50)', 'Nullable', 'RAM (8GB, 16GB...)'],
        ['ScreenSize', 'nvarchar(50)', 'Nullable', 'Kích thước màn hình'],
        ['CPU', 'nvarchar(100)', 'Nullable', 'Vi xử lý'],
        ['CategoryId', 'int', 'FK → Categories', 'Danh mục sản phẩm'],
        ['IsActive', 'bool', 'Default: true', 'Đang hoạt động'],
        ['IsFeatured', 'bool', 'Default: false', 'Sản phẩm nổi bật'],
        ['CreatedAt', 'datetime', 'Default: UTC now', 'Ngày tạo'],
        ['UpdatedAt', 'datetime', 'Nullable', 'Ngày cập nhật gần nhất'],
    ]
)

add_para('BẢNG PRODUCTIMAGES', bold=True, color=NAVY, size=14)
add_table_styled(
    ['Cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['Id', 'int', 'PK, Identity', 'Khóa chính'],
        ['ProductId', 'int', 'FK → Products (Cascade)', 'Liên kết sản phẩm'],
        ['ImagePath', 'nvarchar(500)', 'Required', 'URL ảnh (Cloudinary CDN)'],
        ['SortOrder', 'int', '', 'Thứ tự hiển thị'],
        ['IsMain', 'bool', '', 'Đánh dấu ảnh chính'],
    ]
)

add_para('BẢNG CATEGORIES', bold=True, color=NAVY, size=14)
add_table_styled(
    ['Cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['Id', 'int', 'PK, Identity', 'Khóa chính'],
        ['Name', 'nvarchar(100)', 'Required', 'Tên danh mục'],
        ['Slug', 'nvarchar(100)', 'Required, Unique', 'URL slug'],
        ['Description', 'nvarchar(500)', 'Nullable', 'Mô tả'],
        ['ImagePath', 'nvarchar(500)', 'Nullable', 'Ảnh đại diện'],
        ['SortOrder', 'int', '', 'Thứ tự sắp xếp'],
        ['IsActive', 'bool', 'Default: true', 'Trạng thái'],
        ['CreatedAt', 'datetime', 'Default: UTC now', 'Ngày tạo'],
    ]
)

add_para('BẢNG ASPNETUSERS (kế thừa IdentityUser)', bold=True, color=NAVY, size=14)
add_table_styled(
    ['Cột bổ sung', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['FullName', 'nvarchar(100)', 'Required', 'Họ tên đầy đủ'],
        ['Phone', 'nvarchar(15)', 'Nullable', 'Số điện thoại'],
        ['Address', 'nvarchar(500)', 'Nullable', 'Địa chỉ giao hàng'],
        ['Avatar', 'nvarchar(500)', 'Nullable', 'URL ảnh đại diện'],
        ['CreatedAt', 'datetime', 'Default: UTC now', 'Ngày đăng ký'],
        ['IsActive', 'bool', 'Default: true', 'Tài khoản hoạt động'],
    ]
)
add_para('Các bảng Identity mặc định: AspNetRoles, AspNetUserRoles, AspNetUserClaims, '
         'AspNetUserLogins, AspNetUserTokens, AspNetRoleClaims.', italic=True, color=GRAY, size=11)

doc.add_page_break()

add_para('BẢNG CARTS', bold=True, color=NAVY, size=14)
add_table_styled(
    ['Cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['Id', 'int', 'PK, Identity', 'Khóa chính'],
        ['UserId', 'string', 'FK → AspNetUsers, Unique', 'Mỗi user có 1 giỏ hàng'],
        ['CreatedAt', 'datetime', 'Default: UTC now', 'Ngày tạo'],
        ['UpdatedAt', 'datetime', 'Nullable', 'Lần cập nhật cuối'],
    ]
)

add_para('BẢNG CARTDETAILS', bold=True, color=NAVY, size=14)
add_table_styled(
    ['Cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['Id', 'int', 'PK, Identity', 'Khóa chính'],
        ['CartId', 'int', 'FK → Carts (Cascade)', 'Giỏ hàng chứa'],
        ['ProductId', 'int', 'FK → Products (Restrict)', 'Sản phẩm được chọn'],
        ['Quantity', 'int', 'Range(1, 99)', 'Số lượng'],
        ['UnitPrice', 'decimal(18,0)', '', 'Đơn giá tại thời điểm thêm'],
    ]
)

add_para('BẢNG ORDERS', bold=True, color=NAVY, size=14)
add_table_styled(
    ['Cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['Id', 'int', 'PK, Identity', 'Khóa chính'],
        ['UserId', 'string', 'FK → AspNetUsers (Restrict)', 'Người đặt hàng'],
        ['OrderCode', 'nvarchar(20)', 'Required, Unique', 'Mã đơn (NX-yyyyMMdd-XXXX)'],
        ['FullName', 'nvarchar(100)', 'Required', 'Họ tên người nhận'],
        ['Phone', 'nvarchar(15)', 'Required', 'SĐT người nhận'],
        ['Address', 'nvarchar(500)', 'Required', 'Địa chỉ giao hàng'],
        ['Note', 'nvarchar(500)', 'Nullable', 'Ghi chú đơn hàng'],
        ['TotalAmount', 'decimal(18,0)', '', 'Tổng tiền sau giảm giá'],
        ['DiscountAmount', 'decimal(18,0)', '', 'Số tiền được giảm'],
        ['VoucherCode', 'nvarchar(20)', 'Nullable', 'Mã voucher đã áp dụng'],
        ['Status', 'enum(int)', 'Default: Pending', 'Pending → Confirmed → Shipping → Delivered / Cancelled'],
        ['PaymentMethod', 'nvarchar(50)', 'Default: COD', 'Phương thức thanh toán'],
        ['CreatedAt', 'datetime', 'Default: UTC now', 'Ngày đặt hàng'],
        ['UpdatedAt', 'datetime', 'Nullable', 'Ngày cập nhật trạng thái'],
    ]
)

add_para('BẢNG ORDERDETAILS', bold=True, color=NAVY, size=14)
add_table_styled(
    ['Cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['Id', 'int', 'PK, Identity', 'Khóa chính'],
        ['OrderId', 'int', 'FK → Orders (Cascade)', 'Đơn hàng chứa'],
        ['ProductId', 'int', 'FK → Products (Restrict)', 'Sản phẩm'],
        ['ProductName', 'nvarchar(200)', 'Required', 'Tên SP (snapshot tại thời điểm đặt)'],
        ['Quantity', 'int', 'Range(1, 99)', 'Số lượng đặt'],
        ['UnitPrice', 'decimal(18,0)', '', 'Đơn giá tại thời điểm đặt'],
    ]
)

add_para('BẢNG VOUCHERS', bold=True, color=NAVY, size=14)
add_table_styled(
    ['Cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['Id', 'int', 'PK, Identity', 'Khóa chính'],
        ['Code', 'nvarchar(20)', 'Required, Unique', 'Mã giảm giá (VD: WELCOME10)'],
        ['Description', 'nvarchar(200)', 'Nullable', 'Mô tả voucher'],
        ['DiscountAmount', 'decimal(18,0)', '', 'Giảm cố định (VNĐ)'],
        ['DiscountPercent', 'int', '', 'Giảm theo % (0-100)'],
        ['MinOrderAmount', 'decimal(18,0)', '', 'Giá trị đơn tối thiểu'],
        ['MaxDiscountAmount', 'decimal(18,0)', '', 'Mức giảm tối đa'],
        ['UsageLimit', 'int', '', 'Số lần dùng tối đa'],
        ['UsedCount', 'int', '', 'Số lần đã sử dụng'],
        ['StartDate', 'datetime', '', 'Ngày bắt đầu hiệu lực'],
        ['EndDate', 'datetime', '', 'Ngày hết hạn'],
        ['IsActive', 'bool', 'Default: true', 'Kích hoạt / Vô hiệu hóa'],
    ]
)

add_para('BẢNG SHOPCONFIGS', bold=True, color=NAVY, size=14)
add_table_styled(
    ['Cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['Id', 'int', 'PK, Identity', 'Khóa chính'],
        ['Key', 'nvarchar(100)', 'Required, Unique', 'Tên cấu hình (ShopName, Phone...)'],
        ['Value', 'nvarchar(2000)', 'Nullable', 'Giá trị cấu hình'],
        ['Type', 'nvarchar(20)', 'Default: string', 'Kiểu dữ liệu'],
    ]
)

doc.add_page_break()

# ── 2.3 Giải thích cách thiết kế
add_heading_styled('2.3. Giải thích cách thiết kế', level=2)

add_para('A. Phương pháp Code-First', bold=True, color=NAVY)
add_para(
    'Các bảng được định nghĩa từ C# model class. EF Core tự động tạo migration '
    'và đồng bộ schema với database. Ưu điểm: kiểm soát version CSDL qua code, '
    'dễ dàng rollback và deploy.'
)

add_para('B. Data Annotation + Fluent API', bold=True, color=NAVY)
add_para(
    'Sử dụng kết hợp: Data Annotation ([Required], [MaxLength], [Range], [Column]) '
    'trên model class cho validation. Fluent API trong OnModelCreating() cho cấu hình '
    'quan hệ, index unique, delete behavior (Cascade/Restrict).'
)

add_para('C. Thiết kế quan hệ', bold=True, color=NAVY)
add_table_styled(
    ['Quan hệ', 'Kiểu', 'Delete Behavior', 'Lý do'],
    [
        ['Category → Products', '1:N', 'Restrict', 'Không cho xóa danh mục đang chứa SP'],
        ['Product → ProductImages', '1:N', 'Cascade', 'Xóa SP thì xóa luôn ảnh'],
        ['User → Cart', '1:1', 'Cascade', 'Mỗi user chỉ 1 giỏ, xóa user = xóa giỏ'],
        ['Cart → CartDetails', '1:N', 'Cascade', 'Xóa giỏ = xóa hết items'],
        ['CartDetail → Product', 'N:1', 'Restrict', 'Không cho xóa SP đang trong giỏ'],
        ['User → Orders', '1:N', 'Restrict', 'Không cho xóa user có đơn hàng'],
        ['Order → OrderDetails', '1:N', 'Cascade', 'Xóa đơn = xóa chi tiết đơn'],
        ['OrderDetail → Product', 'N:1', 'Restrict', 'Không cho xóa SP đã bán'],
    ]
)

add_para('D. Thiết kế Snapshot trong OrderDetail', bold=True, color=NAVY)
add_info_box(
    'OrderDetail lưu cả ProductName và UnitPrice tại thời điểm đặt hàng (snapshot), '
    'không chỉ dựa vào FK. Lý do: nếu sản phẩm bị đổi tên hoặc thay đổi giá sau đó, '
    'lịch sử đơn hàng vẫn chính xác với dữ liệu gốc.'
)

add_para('E. ASP.NET Core Identity', bold=True, color=NAVY)
add_para(
    'Hệ thống user kế thừa IdentityUser để tận dụng sẵn: hash mật khẩu (bcrypt), '
    'quản lý role, cookie authentication, lockout policy, token providers. '
    'Bổ sung các trường FullName, Phone, Address, Avatar, IsActive cho nghiệp vụ.'
)

# ── 2.4 Dữ liệu mẫu
add_heading_styled('2.4. Dữ liệu mẫu (Seed Data)', level=2)

add_table_styled(
    ['Loại dữ liệu', 'Số lượng', 'Phương pháp seed', 'Nguồn'],
    [
        ['Sản phẩm', '56', 'Runtime (ProductSeeder.cs)', 'products-seed.json (crawl từ TGDĐ)'],
        ['Danh mục', '8', 'EF HasData (CategorySeed.cs)', 'Định nghĩa trong code'],
        ['Người dùng', '8', 'Runtime (SeedData.cs)', 'Định nghĩa trong code'],
        ['Voucher', '8', 'EF HasData (VoucherSeed.cs)', 'Định nghĩa trong code'],
        ['Cấu hình', '5', 'EF HasData (ShopConfigSeed.cs)', 'Định nghĩa trong code'],
    ]
)

add_para('Tài khoản mẫu:', bold=True, color=NAVY)
add_table_styled(
    ['Email', 'Mật khẩu', 'Vai trò', 'Họ tên'],
    [
        ['admin@nexora.vn', 'Admin@123', 'Admin', 'Admin Nexora'],
        ['staff@nexora.vn', 'Staff@123', 'Staff', 'Nguyễn Văn Bình'],
        ['staff2@nexora.vn', 'Staff@123', 'Staff', 'Trần Thị Cẩm'],
        ['customer@nexora.vn', 'Customer@123', 'Customer', 'Lê Hoàng Dũng'],
        ['ngoclan@gmail.com', 'Customer@123', 'Customer', 'Phạm Ngọc Lan'],
        ['minhtuan@gmail.com', 'Customer@123', 'Customer', 'Võ Minh Tuấn'],
        ['thuhang@gmail.com', 'Customer@123', 'Customer', 'Đặng Thu Hằng'],
        ['quocviet@gmail.com', 'Customer@123', 'Customer', 'Huỳnh Quốc Việt'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 3. SƠ ĐỒ USE CASE
# ══════════════════════════════════════════════════════════════════
add_heading_styled('3. Sơ đồ Use Case', level=1)

add_para('Hệ thống có 3 tác nhân (Actor):', bold=True, color=NAVY)

# Actor descriptions in a nice table
add_table_styled(
    ['Actor', 'Vai trò', 'Mô tả'],
    [
        ['Customer', 'Khách hàng', 'Người dùng đã đăng ký, mua sắm trực tuyến'],
        ['Staff', 'Nhân viên', 'Quản lý sản phẩm, danh mục, duyệt đơn hàng'],
        ['Admin', 'Quản trị viên', 'Toàn quyền hệ thống, quản lý user + cấu hình'],
    ]
)

add_para('')
add_para('USE CASE — KHÁCH HÀNG (Customer)', bold=True, color=NAVY, size=14)
add_table_styled(
    ['Mã', 'Use Case', 'Mô tả chi tiết'],
    [
        ['UC01', 'Đăng ký tài khoản', 'Nhập email, mật khẩu, họ tên, SĐT. Validation đầy đủ, tự gán role Customer.'],
        ['UC02', 'Đăng nhập', 'Email + mật khẩu, cookie auth, Remember Me, lockout sau nhiều lần sai.'],
        ['UC03', 'Đăng xuất', 'Kết thúc phiên, xóa cookie xác thực.'],
        ['UC04', 'Xem danh sách SP', 'Grid sản phẩm, filter theo danh mục, tìm kiếm theo tên, sort giá tăng/giảm.'],
        ['UC05', 'Xem chi tiết SP', 'Gallery ảnh, thông số kỹ thuật (RAM, CPU, Storage, Color), mô tả, giá.'],
        ['UC06', 'Thêm vào giỏ hàng', 'Chọn số lượng, HTMX cập nhật badge. Nếu SP đã có thì cộng dồn.'],
        ['UC07', 'Quản lý giỏ hàng', 'Tăng/giảm số lượng, xóa SP, xem tổng tiền.'],
        ['UC08', 'Đặt hàng', 'Nhập thông tin giao hàng, áp mã giảm giá (Voucher), xác nhận COD.'],
        ['UC09', 'Xem lịch sử đơn', 'Danh sách đơn đã đặt, trạng thái, chi tiết từng đơn.'],
        ['UC10', 'Quản lý hồ sơ', 'Sửa họ tên, SĐT, địa chỉ giao hàng.'],
        ['UC11', 'Đổi mật khẩu', 'Xác thực mật khẩu cũ trước khi đổi.'],
    ]
)

add_para('USE CASE — NHÂN VIÊN (Staff)', bold=True, color=NAVY, size=14)
add_table_styled(
    ['Mã', 'Use Case', 'Mô tả chi tiết'],
    [
        ['UC12', 'Xem Dashboard', 'Thống kê: tổng doanh thu, số đơn, số SP, số KH. Đơn hàng gần đây.'],
        ['UC13', 'CRUD Sản phẩm', 'Thêm/sửa/xóa SP. Upload ảnh Cloudinary. SEO slug tự động.'],
        ['UC14', 'CRUD Danh mục', 'Thêm/sửa/xóa danh mục. Slug auto, sắp xếp, ẩn/hiện.'],
        ['UC15', 'Quản lý đơn hàng', 'Danh sách, cập nhật trạng thái (Pending → Confirmed → Shipping → Delivered).'],
        ['UC16', 'Cập nhật hồ sơ', 'Sửa thông tin cá nhân nhân viên.'],
    ]
)

add_para('USE CASE — QUẢN TRỊ VIÊN (Admin)', bold=True, color=NAVY, size=14)
add_para('Kế thừa tất cả Use Case của Staff, bổ sung:', italic=True)
add_table_styled(
    ['Mã', 'Use Case', 'Mô tả chi tiết'],
    [
        ['UC17', 'Quản lý người dùng', 'Danh sách user, tìm kiếm theo tên/email.'],
        ['UC18', 'Phân quyền', 'Đổi vai trò user (Admin / Staff / Customer). Server-side validation.'],
        ['UC19', 'Khóa/Mở tài khoản', 'Toggle IsActive, ngăn user bị khóa đăng nhập.'],
        ['UC20', 'CRUD Voucher', 'Thêm/sửa/xóa mã giảm giá. Cấu hình %, cố định, min order, max discount, thời hạn.'],
        ['UC21', 'Cấu hình cửa hàng', 'Sửa tên shop, SĐT, email, địa chỉ, mô tả.'],
    ]
)

add_para('')
add_para('SƠ ĐỒ USE CASE TỔNG QUAN', bold=True, color=NAVY, size=14)

uc_diagram = """
                        ┌─────────────────────────────────────────────────────────────┐
                        │                    HỆ THỐNG NEXORA                         │
                        │                                                             │
  ┌───────────┐         │  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐      │
  │           │─────────┼─►│  UC01: Đăng  │  │  UC04: Xem   │  │  UC06: Thêm  │      │
  │  KHÁCH    │         │  │  ký tài khoản│  │  DS sản phẩm │  │  giỏ hàng    │      │
  │  HÀNG     │─────────┼─►│  UC02: Đăng  │  │  UC05: Chi   │  │  UC07: Quản  │      │
  │(Customer) │         │  │  nhập        │  │  tiết SP     │  │  lý giỏ hàng │      │
  │           │─────────┼─►│  UC03: Đăng  │  │  UC08: Đặt   │  │  UC09: Lịch  │      │
  │           │         │  │  xuất        │  │  hàng        │  │  sử đơn hàng │      │
  │           │         │  │  UC10: Hồ sơ │  │  UC11: Đổi   │  │              │      │
  └───────────┘         │  │  cá nhân     │  │  mật khẩu    │  │              │      │
                        │  └──────────────┘  └──────────────┘  └──────────────┘      │
                        │                                                             │
  ┌───────────┐         │  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐      │
  │           │─────────┼─►│  UC12: Xem   │  │  UC13: CRUD  │  │  UC14: CRUD  │      │
  │  NHÂN     │         │  │  Dashboard   │  │  Sản phẩm    │  │  Danh mục    │      │
  │  VIÊN     │─────────┼─►│  UC15: Quản  │  │  UC16: Cập   │  │              │      │
  │  (Staff)  │         │  │  lý đơn hàng │  │  nhật hồ sơ  │  │              │      │
  └───────────┘         │  └──────────────┘  └──────────────┘  └──────────────┘      │
                        │                                                             │
  ┌───────────┐         │  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐      │
  │           │─────────┼─►│  UC17: Quản  │  │  UC18: Phân  │  │  UC19: Khóa/ │      │
  │  ADMIN    │         │  │  lý user     │  │  quyền       │  │  Mở tài khoản│      │
  │           │─────────┼─►│  UC20: Cấu   │  │              │  │              │      │
  │           │         │  │  hình shop   │  │  + Kế thừa   │  │  tất cả UC   │      │
  └───────────┘         │  │              │  │  của Staff   │  │              │      │
                        │  └──────────────┘  └──────────────┘  └──────────────┘      │
                        └─────────────────────────────────────────────────────────────┘"""

add_code_block(uc_diagram, font_size=7)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 4. CẤU TRÚC DỰ ÁN
# ══════════════════════════════════════════════════════════════════
add_heading_styled('4. Cấu trúc dự án ASP.NET Core MVC', level=1)

add_heading_styled('4.1. Mô hình MVC', level=2)

add_para('Ứng dụng tuân theo mô hình Model-View-Controller:', bold=True)

# MVC diagram
mvc_diagram = """
  ┌─────────────┐     HTTP Request      ┌──────────────────┐
  │             │ ───────────────────►   │                  │
  │   Browser   │                        │   Controller     │
  │   (Client)  │   ◄───────────────── │   (C# classes)   │
  │             │     HTTP Response      │                  │
  └─────────────┘                        └────┬────────┬────┘
                                              │        │
                                    Query DB  │        │  Select View
                                              ▼        ▼
                                   ┌──────────┐  ┌──────────┐
                                   │  Model   │  │   View   │
                                   │ (EF Core │  │ (Razor   │
                                   │  DbSet)  │  │  .cshtml)│
                                   └──────────┘  └──────────┘"""

add_code_block(mvc_diagram, font_size=9)

add_bullet('Định nghĩa cấu trúc dữ liệu (Product, Category, Order...) và business logic. '
           'Sử dụng Data Annotation cho validation.', bold_prefix='Model: ')
add_bullet('Razor View (.cshtml) kết hợp Tailwind CSS + DaisyUI. Sử dụng Layout chung '
           '(_Layout.cshtml cho storefront, _AdminLayout.cshtml cho admin panel).', bold_prefix='View: ')
add_bullet('Xử lý HTTP request, tương tác với database qua EF Core, trả về View hoặc redirect. '
           'Dependency Injection cho DbContext, UserManager, CloudinaryService.', bold_prefix='Controller: ')

add_heading_styled('4.2. Cấu trúc thư mục', level=2)

structure = """Nexora/
├── Controllers/                       # Xử lý request
│   ├── HomeController.cs              # Trang chủ (featured, new arrivals)
│   ├── ProductController.cs           # Danh sách, chi tiết, tìm kiếm SP
│   ├── CartController.cs              # Giỏ hàng (Add, Update, Remove, Count)
│   ├── OrderController.cs             # Checkout, lịch sử, chi tiết đơn
│   ├── AccountController.cs           # Đăng ký, đăng nhập, hồ sơ, đổi MK
│   ├── AdminController.cs             # Dashboard thống kê
│   ├── AdminProductController.cs      # CRUD SP + upload Cloudinary
│   ├── AdminCategoryController.cs     # CRUD danh mục
│   ├── AdminOrderController.cs        # Quản lý + cập nhật trạng thái đơn
│   ├── AdminUserController.cs         # Quản lý user + phân quyền
│   ├── AdminVoucherController.cs      # CRUD voucher/mã giảm giá
│   └── AdminConfigController.cs       # Cấu hình cửa hàng
├── Models/                            # Entities + ViewModels
│   ├── Product.cs, Category.cs, Cart.cs, Order.cs, Voucher.cs...
│   └── ViewModels/
│       ├── LoginViewModel.cs, RegisterViewModel.cs
│       ├── ProfileViewModel.cs, CheckoutViewModel.cs
│       └── ProductFormViewModel.cs
├── Views/                             # Razor Views
│   ├── Shared/_Layout.cshtml          # Layout khách hàng
│   ├── Shared/_AdminLayout.cshtml     # Layout admin
│   ├── Home/, Product/, Cart/, Order/, Account/
│   └── Admin*/ (Product, Category, Order, User, Voucher, Config)
├── Data/                              # Database layer
│   ├── ApplicationDbContext.cs        # DbContext + Fluent API config
│   ├── SeedData.cs                    # Seed users & roles (runtime)
│   ├── ProductSeeder.cs               # Seed 56 SP từ JSON (runtime)
│   ├── products-seed.json             # Dữ liệu 56 sản phẩm
│   └── Seeds/                         # EF HasData seeds
│       ├── CategorySeed.cs, VoucherSeed.cs, ShopConfigSeed.cs
├── Services/
│   └── CloudinaryService.cs           # Upload/delete ảnh Cloudinary
├── Migrations/                        # EF Core migrations
├── wwwroot/                           # Static files (CSS, JS, images)
└── Program.cs                         # Entry point + DI + middleware"""

add_code_block(structure, font_size=8)

add_heading_styled('4.3. Công nghệ sử dụng', level=2)

add_table_styled(
    ['Công nghệ', 'Phiên bản', 'Vai trò trong dự án'],
    [
        ['ASP.NET Core', '8.0', 'Web framework chính (MVC pattern)'],
        ['Entity Framework Core', '8.0', 'ORM — Code-First, Migration, LINQ to Entities'],
        ['ASP.NET Core Identity', '8.0', 'Xác thực, phân quyền, quản lý user/role'],
        ['PostgreSQL', '16+', 'CSDL quan hệ (hosted trên Neon cloud)'],
        ['Npgsql', '8.0', 'PostgreSQL driver cho .NET'],
        ['CloudinaryDotNet', '1.26+', 'Upload, lưu trữ, tối ưu ảnh trên CDN'],
        ['Tailwind CSS', '3.x', 'CSS utility-first, responsive design'],
        ['DaisyUI', '4.x', 'Component library trên Tailwind (table, card, badge...)'],
        ['HTMX', '1.9+', 'AJAX partial update (cập nhật giỏ hàng không reload)'],
    ]
)

add_para('Kỹ thuật chính:', bold=True, color=NAVY)
add_bullet('Định nghĩa service trong Program.cs, inject qua constructor. '
           'VD: DbContext, UserManager, SignInManager, CloudinaryService.', bold_prefix='Dependency Injection: ')
add_bullet('Truy vấn dữ liệu type-safe bằng C# lambda. EF Core dịch sang SQL. '
           'VD: .Include().Where().OrderBy().ToListAsync().', bold_prefix='LINQ to Entities: ')
add_bullet('[Required], [MaxLength], [Range], [Column] cho validation + DB mapping.', bold_prefix='Data Annotation: ')
add_bullet('Cấu hình quan hệ, index, delete behavior trong OnModelCreating().', bold_prefix='Fluent API: ')
add_bullet('Cookie authentication qua ASP.NET Identity, tự động quản lý session.', bold_prefix='Session/Cookie: ')
add_bullet('asp-action, asp-controller, asp-for, asp-validation-for trong Razor.', bold_prefix='Tag Helpers: ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 5. CHỨC NĂNG ĐÃ THỰC HIỆN
# ══════════════════════════════════════════════════════════════════
add_heading_styled('5. Các chức năng đã thực hiện', level=1)

add_heading_styled('5.1. Chức năng khách hàng', level=2)

add_table_styled(
    ['Chức năng', 'Chi tiết', 'URL'],
    [
        ['Trang chủ', 'SP nổi bật, mới nhất, phân theo danh mục', '/'],
        ['Danh sách SP', 'Filter danh mục, tìm kiếm tên, sort giá', '/Product'],
        ['Chi tiết SP', 'Gallery ảnh, specs, mô tả, giá', '/Product/Detail?slug=...'],
        ['Đăng ký', 'Email, MK, họ tên, SĐT. Auto role Customer', '/Account/Register'],
        ['Đăng nhập', 'Cookie auth, Remember Me, Lockout', '/Account/Login'],
        ['Giỏ hàng', 'Thêm/sửa SL/xóa SP. HTMX badge update', '/Cart'],
        ['Đặt hàng', 'Form giao hàng, áp Voucher, COD', '/Order/Checkout'],
        ['Lịch sử đơn', 'Danh sách đơn, chi tiết từng đơn', '/Order/MyOrders'],
        ['Hồ sơ cá nhân', 'Sửa tên, SĐT, địa chỉ', '/Account/Profile'],
        ['Đổi mật khẩu', 'Xác thực MK cũ trước khi đổi', '/Account/ChangePassword'],
    ]
)

add_heading_styled('5.2. Chức năng quản lý (Admin/Staff)', level=2)

add_table_styled(
    ['Chức năng', 'Chi tiết', 'Quyền'],
    [
        ['Dashboard', 'Doanh thu, đơn hàng, SP, KH. Đơn gần đây.', 'Admin, Staff'],
        ['CRUD Sản phẩm', 'Thêm/sửa/xóa. Upload nhiều ảnh qua Cloudinary hoặc nhập URL. SEO slug tự động.', 'Admin, Staff'],
        ['CRUD Danh mục', 'Thêm/sửa/xóa. Auto slug, sort order, ẩn/hiện.', 'Admin, Staff'],
        ['Quản lý đơn hàng', 'Danh sách, đổi trạng thái (5 bước).', 'Admin, Staff'],
        ['CRUD Voucher', 'Thêm/sửa/xóa mã giảm giá. Giảm %, cố định, đơn tối thiểu, giảm tối đa, thời hạn.', 'Admin only'],
        ['Quản lý người dùng', 'Danh sách, tìm kiếm, đổi role, khóa/mở.', 'Admin only'],
        ['Cấu hình cửa hàng', 'Tên shop, SĐT, email, địa chỉ, mô tả.', 'Admin only'],
    ]
)

add_heading_styled('5.3. Chức năng nâng cao', level=2)

add_bullet('Hỗ trợ giảm theo phần trăm (%) hoặc số tiền cố định (VNĐ). '
           'Có giới hạn: đơn tối thiểu, giảm tối đa, số lần sử dụng, thời hạn hiệu lực. '
           'Validation trên server.', bold_prefix='Voucher/Mã giảm giá: ')
add_bullet('Sản phẩm hỗ trợ upload nhiều file ảnh cùng lúc lên Cloudinary CDN. '
           'Tự động tối ưu chất lượng (auto quality) và định dạng (auto format). '
           'Fallback: nhập URL ảnh trực tiếp.', bold_prefix='Upload ảnh Cloudinary: ')
add_bullet('Slug được tạo tự động từ tên sản phẩm/danh mục (bỏ dấu, lowercase, thay space bằng dash). '
           'Chống trùng slug bằng cách thêm suffix timestamp.', bold_prefix='SEO-friendly URL: ')
add_bullet('Giao diện responsive trên mobile (375px), tablet (768px), desktop (1440px). '
           'Sử dụng Tailwind CSS breakpoints.', bold_prefix='Responsive Design: ')
add_bullet('Cập nhật số lượng giỏ hàng trên navbar mà không reload trang (HTMX request).', bold_prefix='HTMX Partial Update: ')
add_bullet('[Authorize(Roles = "Admin,Staff")] trên controller. Server-side whitelist khi đổi role.', bold_prefix='Phân quyền 3 cấp: ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 6. BẢO MẬT VÀ PHÂN QUYỀN
# ══════════════════════════════════════════════════════════════════
add_heading_styled('6. Bảo mật và phân quyền', level=1)

add_para('Ma trận phân quyền:', bold=True, color=NAVY)
add_table_styled(
    ['Chức năng', 'Customer', 'Staff', 'Admin'],
    [
        ['Storefront (xem SP, tìm kiếm)', 'O', 'O', 'O'],
        ['Giỏ hàng, đặt hàng', 'O', 'O', 'O'],
        ['Hồ sơ cá nhân, đổi MK', 'O', 'O', 'O'],
        ['Dashboard thống kê', 'X', 'O', 'O'],
        ['CRUD Sản phẩm + Upload ảnh', 'X', 'O', 'O'],
        ['CRUD Danh mục', 'X', 'O', 'O'],
        ['Quản lý đơn hàng (admin)', 'X', 'O', 'O'],
        ['CRUD Voucher', 'X', 'X', 'O'],
        ['Quản lý người dùng', 'X', 'X', 'O'],
        ['Phân quyền / Khóa TK', 'X', 'X', 'O'],
        ['Cấu hình cửa hàng', 'X', 'X', 'O'],
    ]
)
add_para('O = Có quyền truy cập  |  X = Bị chặn (redirect AccessDenied hoặc Login)', italic=True, color=GRAY, size=11)

add_para('')
add_para('Các biện pháp bảo mật:', bold=True, color=NAVY)

security_items = [
    ('CSRF Protection', 'Tất cả form POST có [ValidateAntiForgeryToken] + anti-forgery token. Ngăn chặn cross-site request forgery.'),
    ('XSS Prevention', 'Razor View tự động HTML-encode output. Không sử dụng Html.Raw() cho user-generated content.'),
    ('Account Lockout', 'Tài khoản tự động bị khóa tạm thời sau nhiều lần đăng nhập sai (lockoutOnFailure: true).'),
    ('Role Validation', 'Server-side whitelist ["Admin", "Staff", "Customer"] kiểm tra role hợp lệ trước khi thay đổi. Ngăn injection role tùy ý.'),
    ('Password Policy', 'Tối thiểu 6 ký tự, yêu cầu chữ hoa, chữ số, ký tự đặc biệt. Hash bằng Identity (bcrypt).'),
    ('Credential Security', 'Connection string và API key không nằm trong source code. Local: appsettings.Development.json (gitignored). Production: environment variables (Railway).'),
]
add_table_styled(
    ['Biện pháp', 'Chi tiết'],
    security_items
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 7. TRIỂN KHAI
# ══════════════════════════════════════════════════════════════════
add_heading_styled('7. Triển khai (Deployment)', level=1)

deploy_diagram = """
  ┌─────────────────────────────────────────────────────────────┐
  │                    KIẾN TRÚC TRIỂN KHAI                     │
  │                                                             │
  │  ┌─────────────┐     ┌─────────────┐     ┌──────────────┐  │
  │  │   Browser   │────►│  Railway    │────►│  PostgreSQL  │  │
  │  │   (Client)  │◄────│  (Hosting)  │◄────│  (Neon DB)   │  │
  │  └─────────────┘     └──────┬──────┘     └──────────────┘  │
  │                             │                               │
  │                             │ Upload ảnh                    │
  │                             ▼                               │
  │                      ┌─────────────┐                        │
  │                      │ Cloudinary  │                        │
  │                      │ (Image CDN) │                        │
  │                      └─────────────┘                        │
  └─────────────────────────────────────────────────────────────┘"""

add_code_block(deploy_diagram, font_size=9)

add_para('')
add_table_styled(
    ['Thành phần', 'Công nghệ', 'Môi trường'],
    [
        ['Web Server', 'ASP.NET Core 8 (Kestrel)', 'Railway (production) / localhost (dev)'],
        ['Database', 'PostgreSQL 16', 'Neon cloud (cả dev & prod)'],
        ['Image Storage', 'Cloudinary CDN', 'Cloud (auto quality + format)'],
        ['CI/CD', 'Railway auto-deploy', 'Push to main → auto deploy'],
    ]
)

add_para('')
add_para('Cách chạy ứng dụng tại local:', bold=True, color=NAVY)

steps = [
    ('1.', 'Clone repository: git clone https://github.com/kiettt23/nexora.git'),
    ('2.', 'Tạo file appsettings.Development.json với connection string PostgreSQL và Cloudinary keys'),
    ('3.', 'Chạy ứng dụng: dotnet run'),
    ('4.', 'Truy cập: http://localhost:5215'),
    ('5.', 'Database tự động migrate + seed data khi khởi động (56 SP, 8 users, 8 vouchers)'),
]
for num, step in steps:
    p = doc.add_paragraph()
    r1 = p.add_run(num + ' ')
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(13)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(step)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(13)
    r2.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)

add_para('')
add_info_box(
    'Lưu ý: File appsettings.json chỉ chứa template (giá trị rỗng). '
    'Credentials thật nằm trong appsettings.Development.json (đã gitignore) '
    'hoặc environment variables trên Railway (production).'
)

# ══════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'BaoCao.docx')
doc.save(output_path)
print(f'OK: {output_path}')
print(f'Size: {os.path.getsize(output_path) / 1024:.0f} KB')
